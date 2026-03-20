# -*- coding: utf-8 -*-
"""
تحليل السيرة الذاتية وتوليد رسالة تغطية احترافية ومختصرة وصادقة.
يعتمد التقديم التلقائي على GEMINI_API_KEY (من config) لتوليد الرسائل وقراءة السيرة من الصور.
"""
import os
import io
from dotenv import load_dotenv

load_dotenv()

def _gemini_api_key() -> str:
    """مفتاح جيميني من config أو من متغير البيئة (للتقديم التلقائي وقراءة الصور)."""
    try:
        from config import GEMINI_API_KEY
        return (GEMINI_API_KEY or "").strip()
    except ImportError:
        return os.getenv("GEMINI_API_KEY", "").strip()


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """استخراج النص من ملف PDF."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text.strip()
    except ImportError:
        pass
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text.strip()
    except Exception:
        pass
    return ""


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """استخراج النص من ملف Word (DOCX)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_text_from_image(image_bytes: bytes, mime_type: str = "image/jpeg") -> str:
    """استخراج النص من صورة سيرة ذاتية باستخدام Gemini (رؤية) — يعتمد على GEMINI_API_KEY."""
    api_key = _gemini_api_key()
    if not api_key or len(image_bytes) < 50:
        return ""
    try:
        from PIL import Image
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash")
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode == "RGBA":
            img = img.convert("RGB")
        prompt = (
            "استخرج كل النص الظاهر في هذه الصورة (سيرة ذاتية / CV) كما هو، "
            "سطراً سطراً، بدون تفسير أو تعليق. أعد النص فقط."
        )
        response = model.generate_content([img, prompt])
        if response and response.text:
            return response.text.strip()
    except Exception:
        pass
    return ""


def _pdf_page_to_image_bytes(pdf_bytes: bytes, page_index: int = 0, dpi: int = 150) -> bytes | None:
    """تحويل صفحة من PDF إلى صورة (PNG) لقراءتها عبر جيميني — للملفات الممسوحة ضوئياً."""
    try:
        import fitz
        from PIL import Image
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if page_index >= len(doc):
            doc.close()
            return None
        page = doc[page_index]
        pix = page.get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        doc.close()
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def extract_text_from_cv(file_bytes: bytes, filename: str) -> str:
    """
    يقرأ السيرة الذاتية من:
    - ملفات PDF: استخراج النص مباشرة؛ إن كان الملف ممسوحاً ضوئياً (لا نص) يُحوّل لصورة وتُقرأ عبر جيميني.
    - ملفات Word (DOCX): استخراج النص.
    - الصور (JPG/PNG/WebP...): قراءة عبر جيميني (رؤية).
    """
    name = (filename or "").lower().strip()
    if not name and not file_bytes:
        return ""

    if name.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
        # إذا لم يُستخرج نص كافٍ (ملف ممسوح ضوئياً)، نحوّل الصفحة الأولى لصورة ونقرأها بجيميني
        if (not text or len(text.strip()) < 80) and len(file_bytes) > 100:
            img_bytes = _pdf_page_to_image_bytes(file_bytes)
            if img_bytes:
                text = extract_text_from_image(img_bytes, "image/png")
        return text or ""

    if name.endswith(".docx") or name.endswith(".doc"):
        return extract_text_from_docx(file_bytes)

    if any(name.endswith(ext) for ext in (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp")):
        mime = "image/jpeg"
        if name.endswith(".png"):
            mime = "image/png"
        elif name.endswith(".webp"):
            mime = "image/webp"
        elif name.endswith(".gif"):
            mime = "image/gif"
        elif name.endswith(".bmp"):
            mime = "image/bmp"
        return extract_text_from_image(file_bytes, mime)

    return ""


def analyze_cv_and_generate_letter(
    job_title: str,
    applicant_name: str,
    company: str = "",
    job_description: str = "",
    lang: str = "ar",
    cv_text: str = "",
) -> str:
    """
    يحلل السيرة الذاتية ويكتب رسالة تغطية مختصرة وصادقة ومهنية.
    - يستخرج المهارات والخبرات الحقيقية من السيرة
    - لا يدّعي خبرات غير موجودة
    - إذا لا توجد خبرة مباشرة يبرز المهارات القابلة للنقل
    - الرسالة قصيرة: فقرة إلى فقرتين كحد أقصى
    - التقديم التلقائي مربوط بـ GEMINI_API_KEY (config).
    """
    api_key = _gemini_api_key()
    if not api_key:
        return _fallback_cover(job_title, applicant_name, company, lang)

    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash")

        has_cv = cv_text and len(cv_text.strip()) > 80

        if lang == "ar":
            if has_cv:
                job_desc_block = f"""
**وصف الوظيفة ومتطلباتها (اقرأها جيداً وربطها بالسيرة):**
{job_description[:1200] if job_description else "—"}
""" if job_description else ""
                prompt = f"""أنت خبير في كتابة رسائل التقديم الوظيفي الاحترافية.

مهمتك: اكتب رسالة تقديم **مبنية على قراءة الوظيفة ووصفها ثم السيرة الذاتية**، بحيث يظهر التطابق بين المتطلبات والمؤهلات بدون تملق أو كلام فارغ.

**الوظيفة:**
- المسمى: {job_title}
{"- الشركة: " + company if company else ""}
{job_desc_block}

**المتقدم:** {applicant_name}

**نص السيرة الذاتية (استخدمه لربط مؤهلاته بالوظيفة):**
{cv_text[:2500]}

**قواعد صارمة:**
1. اقرأ وصف الوظيفة أولاً وحدد المتطلبات أو المهارات المطلوبة
2. اقرأ السيرة واختر ما يطابق الوظيفة فعلياً (مهارة، خبرة، مؤهل) واذكره بوضوح في الرسالة
3. لا تختلق أو تضخم خبرات غير موجودة في السيرة
4. أسلوب احترافي ومباشر فقط — **ممنوع التملق أو المدح الفارغ** (مثل "شركتكم الرائدة"، "أنا الأفضل"). ركّز على الحقائق والمؤهلات
5. فقرة واحدة أو فقرتين كحد أقصى، بدون عناوين أو نقاط
6. ابدأ بـ "السلام عليكم" وانهِ بالشكر والاسم فقط
7. اذكر **على الأقل نقطتين محددتين من السيرة** (مثل سنوات خبرة/أدوات/إنجاز/دور سابق) بصياغة طبيعية داخل الفقرة
8. لا تستخدم عبارات عامة مبهمة مثل "لدي خبرة واسعة" بدون دليل من السيرة

اكتب الرسالة الآن:"""
            else:
                prompt = f"""اكتب رسالة تقديم مهنية ومختصرة بالعربية لشخص يتقدم لوظيفة:
- المسمى: {job_title}
{"- الشركة: " + company if company else ""}
- اسم المتقدم: {applicant_name}

القواعد: فقرة واحدة، مهنية ومباشرة، بدون تملق أو مدح فارغ. ابدأ بـ "السلام عليكم" وانهِ بالاسم. بدون عناوين أو نقاط.
ممنوع الجمل العامة المكررة."""
        else:
            if has_cv:
                job_desc_block = f"""
**Job description and requirements (read carefully and match to CV):**
{job_description[:1200] if job_description else "—"}
""" if job_description else ""
                prompt = f"""You are an expert professional cover letter writer.

Task: Write a cover letter **based on reading the job and its description first, then the CV**, so the match between requirements and qualifications is clear. Professional and substantive — no flattery or empty praise.

**Job:**
- Title: {job_title}
{"- Company: " + company if company else ""}
{job_desc_block}

**Applicant:** {applicant_name}

**CV content (use it to link his/her qualifications to the job):**
{cv_text[:2500]}

**Strict rules:**
1. Read the job description first and identify key requirements or skills
2. Read the CV and pick what actually matches the job; state it clearly in the letter
3. Never fabricate or exaggerate experience
4. Professional, direct tone only — **no flattery or empty praise** (e.g. "your leading company", "I am the best"). Stick to facts and qualifications
5. Maximum 1–2 short paragraphs, no headers or bullets
6. Start with "Dear Hiring Manager," end with thanks and name only
7. Mention **at least 2 concrete CV facts** (tools, years, responsibilities, outcomes) naturally in the letter
8. Avoid vague statements like "I have extensive experience" without evidence from the CV

Write the letter now:"""
            else:
                prompt = f"""Write a short professional cover letter in English for: Job {job_title}{", " + company if company else ""}. Applicant: {applicant_name}. One paragraph, professional and direct, no flattery. Start with "Dear Hiring Manager," end with name. No formatting."""

        response = model.generate_content(prompt)
        if response and response.text:
            return response.text.strip()
    except Exception:
        pass

    return _fallback_cover(job_title, applicant_name, company, lang)


def generate_cover_letter(
    job_title: str,
    applicant_name: str,
    company: str = "",
    job_description_snippet: str = "",
    lang: str = "ar",
    cv_text: str = "",
) -> str:
    """واجهة موحدة لتوليد رسالة التغطية."""
    return analyze_cv_and_generate_letter(
        job_title=job_title,
        applicant_name=applicant_name,
        company=company,
        job_description=job_description_snippet,
        lang=lang,
        cv_text=cv_text,
    )


def _fallback_cover(job_title: str, applicant_name: str, company: str = "", lang: str = "ar") -> str:
    if lang == "en":
        co = f" at {company}" if company else ""
        return (
            f"Dear Hiring Manager,\n\n"
            f"I am writing to express my interest in the {job_title} position{co}. "
            f"I believe my background makes me a suitable candidate for this role. "
            f"I would welcome the opportunity to discuss my qualifications further.\n\n"
            f"Thank you for your consideration.\n{applicant_name}"
        )
    co = f" في {company}" if company else ""
    return (
        f"السلام عليكم،\n\n"
        f"أتقدم بطلبي للنظر في وظيفة «{job_title}»{co}. "
        f"أثق بأن مؤهلاتي وخبراتي تجعلني مرشحاً مناسباً لهذا الدور، "
        f"وأتطلع لمناقشة كيف يمكنني إضافة قيمة لفريقكم.\n\n"
        f"شكراً لكم،\n{applicant_name}"
    )
